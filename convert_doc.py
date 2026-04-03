#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys

try:
    from docx import Document
    
    doc_path = r'c:\Users\admin\IWE\DS-strategy\inbox\Еже ИФ текст лекций (2025).doc'
    output_path = r'c:\Users\admin\IWE\DS-strategy\inbox\Еже-лекции.txt'
    
    doc = Document(doc_path)
    text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    print(f"✓ Converted to: {output_path}")
    print(f"✓ Paragraphs: {len(doc.paragraphs)}")
    sys.exit(0)

except ImportError:
    print("python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)
except Exception as e:
    print(f"Error: {e}")
    sys.exit(1)
